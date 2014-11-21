# encoding: utf-8
import re
from docx.api import Document
from docx.table import _Cell
from docx.text import Paragraph


class TagDispatcher(object):
    def __init__(self):
        super(TagDispatcher, self).__init__()

    @classmethod
    def append_head(cls, element, container):
        raise NotImplementedError("Implemented in inheriting classes")

    @classmethod
    def append_tail(cls, element, container):
        raise NotImplementedError("Implemented in inheriting classes")

    @classmethod
    def get_current_paragraph(cls, container):
        current_paragraph = container
        if isinstance(container, Paragraph):
            if isinstance(container._parent, _Cell):
                current_paragraph = container._parent.paragraphs[-1]

        if isinstance(container, Document):
            current_paragraph = container.add_paragraph()
        return current_paragraph

    @classmethod
    def get_new_paragraph(cls, container):
        new_paragraph = container
        if isinstance(container, Paragraph):
            if isinstance(container._parent, _Cell):
                new_paragraph = container._parent.paragraphs[0]
                if len(container._parent.paragraphs) > 1:
                    new_paragraph = container._parent.add_paragraph()
                else:
                    if container._parent.paragraphs[0].text:
                        new_paragraph = container._parent.add_paragraph()
            else:
                if container.text:
                    new_paragraph = container._parent.add_paragraph()
        if isinstance(container, Document):
            new_paragraph = container.add_paragraph()
        return new_paragraph


def replace_whitespaces(text):
    """
    replaces multiple whitespaces and line breaks by a single whitespace
    """
    if text:
        text = ' '.join(text.split('\n'))
        text = re.sub(' +', ' ', text)

    return text if text else ''